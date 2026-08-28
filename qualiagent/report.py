"""Build DOCX analysis reports from persisted run data."""

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document
from docx.document import Document as DocxDocument
from sqlalchemy import func, select
from sqlalchemy.orm import Session, selectinload

from qualiagent.models import AnalysisRun, Citation, Section, Source, Study
from qualiagent.respondent_counts import count_study_respondents


@dataclass
class CitationTableRow:
    """One row in the report citation table."""

    marker: str
    respondent: str
    source_code: str
    quoted_text: str
    verified: bool


@dataclass
class MethodologySummary:
    """Counts and flags for the methodological note."""

    source_count: int
    respondent_count: int
    thin_questions: list[str]
    unverified_citation_count: int


def load_analysis_run_with_sections(session: Session, analysis_run_id: UUID) -> AnalysisRun:
    """Load an analysis run with ordered sections and citations.

    Args:
        session: Database session.
        analysis_run_id: Analysis run UUID.

    Returns:
        Analysis run with relationships loaded.

    Raises:
        ValueError: If the run does not exist.
    """
    statement = (
        select(AnalysisRun)
        .where(AnalysisRun.id == analysis_run_id)
        .options(selectinload(AnalysisRun.sections).selectinload(Section.citations))
    )
    analysis_run = session.execute(statement).scalar_one_or_none()
    if analysis_run is None:
        raise ValueError(f"AnalysisRun {analysis_run_id} not found")
    return analysis_run


def build_citation_table_rows(session: Session, sections: list[Section]) -> list[CitationTableRow]:
    """Build citation table rows with respondent and source metadata.

    Args:
        session: Database session.
        sections: Sections belonging to one analysis run.

    Returns:
        Ordered citation rows for the DOCX table.
    """
    citations: list[Citation] = []
    for section in sorted(sections, key=lambda item: item.position):
        citations.extend(section.citations)

    if not citations:
        return []

    source_ids = {citation.source_id for citation in citations}
    sources = session.execute(select(Source).where(Source.id.in_(source_ids))).scalars().all()
    source_by_id = {source.id: source for source in sources}

    rows: list[CitationTableRow] = []
    for citation in citations:
        source = source_by_id.get(citation.source_id)
        respondent = "—"
        source_code = "—"
        if source is not None:
            respondent = source.respondent_label or f"source:{source.source_code}"
            source_code = source.source_code
        rows.append(
            CitationTableRow(
                marker=citation.marker,
                respondent=respondent,
                source_code=source_code,
                quoted_text=citation.quoted_text,
                verified=citation.verified,
            )
        )
    return rows


def build_methodology_summary(
    session: Session,
    study_id: UUID,
    sections: list[Section],
) -> MethodologySummary:
    """Compute methodological note values for a study run.

    Args:
        session: Database session.
        study_id: Study UUID.
        sections: Sections from the analysis run.

    Returns:
        Methodology summary used in the DOCX note.
    """
    source_count = session.execute(
        select(func.count()).select_from(Source).where(Source.study_id == study_id)
    ).scalar_one()
    thin_questions = [
        section.research_question
        for section in sorted(sections, key=lambda item: item.position)
        if section.coverage == "thin"
    ]
    unverified_citation_count = sum(
        1 for section in sections for citation in section.citations if not citation.verified
    )
    return MethodologySummary(
        source_count=int(source_count),
        respondent_count=count_study_respondents(session, study_id),
        thin_questions=thin_questions,
        unverified_citation_count=unverified_citation_count,
    )


def fill_report_document(
    document: DocxDocument,
    study_name: str,
    sections: list[Section],
    citation_rows: list[CitationTableRow],
    methodology: MethodologySummary,
) -> None:
    """Write report content into an open Document.

    Args:
        document: Target python-docx document.
        study_name: Study title.
        sections: Ordered report sections.
        citation_rows: Citation table rows.
        methodology: Methodological note values.
    """
    document.add_heading(study_name, level=0)
    document.add_paragraph("Raport z analizy jakościowej QualiAgent.")

    document.add_heading("Rozdziały", level=1)
    if not sections:
        document.add_paragraph("Brak sekcji w tym przebiegu analizy.")
    for index, section in enumerate(sorted(sections, key=lambda item: item.position), start=1):
        document.add_heading(f"{index}. {section.research_question}", level=2)
        document.add_paragraph(section.body)
        document.add_paragraph(
            f"Pokrycie: {section.coverage}. "
            f"Respondenci: {section.respondents_covered}/{section.respondents_total}. "
            f"{section.coverage_note}".strip()
        )

    document.add_heading("Tabela cytowań", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("Marker", "Respondent", "Źródło", "Fragment", "Weryfikacja")
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
    if not citation_rows:
        empty = table.add_row().cells
        empty[0].text = "—"
        empty[1].text = "—"
        empty[2].text = "—"
        empty[3].text = "Brak cytowań."
        empty[4].text = "—"
    for row in citation_rows:
        cells = table.add_row().cells
        cells[0].text = row.marker
        cells[1].text = row.respondent
        cells[2].text = row.source_code
        cells[3].text = row.quoted_text
        cells[4].text = "zweryfikowany" if row.verified else "niezweryfikowany"

    document.add_heading("Nota metodologiczna", level=1)
    document.add_paragraph(f"Liczba źródeł: {methodology.source_count}.")
    document.add_paragraph(f"Liczba respondentów: {methodology.respondent_count}.")
    if methodology.thin_questions:
        document.add_paragraph("Pytania z cienkim pokryciem:")
        for question in methodology.thin_questions:
            document.add_paragraph(question, style="List Bullet")
    else:
        document.add_paragraph("Pytania z cienkim pokryciem: brak.")
    document.add_paragraph(f"Liczba cytatów niezweryfikowanych: {methodology.unverified_citation_count}.")


def build_analysis_run_report(session: Session, analysis_run_id: UUID) -> bytes:
    """Build a DOCX report for one analysis run.

    Args:
        session: Database session.
        analysis_run_id: Analysis run UUID.

    Returns:
        DOCX file bytes.
    """
    analysis_run = load_analysis_run_with_sections(session, analysis_run_id)
    study = session.get(Study, analysis_run.study_id)
    if study is None:
        raise ValueError(f"Study {analysis_run.study_id} not found")

    sections = list(analysis_run.sections)
    citation_rows = build_citation_table_rows(session, sections)
    methodology = build_methodology_summary(session, study.id, sections)

    document = Document()
    fill_report_document(
        document=document,
        study_name=study.name,
        sections=sections,
        citation_rows=citation_rows,
        methodology=methodology,
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_analysis_run_report(
    session: Session,
    analysis_run_id: UUID,
    output_path: Path,
) -> Path:
    """Build and write a DOCX report to disk.

    Args:
        session: Database session.
        analysis_run_id: Analysis run UUID.
        output_path: Destination ``.docx`` path.

    Returns:
        The written path.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(build_analysis_run_report(session, analysis_run_id))
    return output_path
