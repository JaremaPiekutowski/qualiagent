from pathlib import Path

from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session

from qualiagent.config import Settings
from qualiagent.ingest.source_ingest import ingest_source
from qualiagent.models import Chunk, Source, Study
from tests.stub_embedding import StubEmbeddingClient


def write_sample_txt(directory: Path) -> Path:
    path = directory / "interview_r01.txt"
    path.write_text(
        "Zmiana organizacyjna była odgórna.\n\n"
        "Nikt nas nie pytał o zdanie przed wdrożeniem nowego procesu.\n\n"
        "Po miesiącu praca stała się bardziej chaotyczna niż wcześniej.",
        encoding="utf-8",
    )
    return path


def write_sample_docx(directory: Path) -> Path:
    path = directory / "interview_r02.docx"
    document = Document()
    document.add_paragraph("Dla mnie zmiana była okazją do nauki.")
    document.add_paragraph("Szefowie komunikowali cele jasno, ale brakowało wsparcia.")
    document.add_paragraph("Najtrudniejszy był pierwszy tydzień po migracji systemów.")
    document.save(path)
    return path


def write_sample_pdf(directory: Path) -> Path:
    # Minimal valid-enough PDF with extractable text for pdfminer.
    path = directory / "interview_r03.pdf"
    content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 68 >>
stream
BT /F1 12 Tf 50 100 Td (Zmiana byla stresujaca dla zespolu.) Tj ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000385 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
462
%%EOF
"""
    path.write_bytes(content)
    return path


def test_ingest_three_files_creates_chunks_with_embeddings(
    session: Session,
    settings: Settings,
    embedding_client: StubEmbeddingClient,
    tmp_path: Path,
) -> None:
    study = Study(
        name="Zmiana organizacyjna",
        research_questions=["Jak respondenci opisują zmianę?"],
        web_search_enabled=False,
    )
    session.add(study)
    session.flush()

    files = [
        (write_sample_txt(tmp_path), "R01"),
        (write_sample_docx(tmp_path), "R02"),
        (write_sample_pdf(tmp_path), "R03"),
    ]

    sources: list[Source] = []
    for file_path, respondent_label in files:
        source = ingest_source(
            session=session,
            study_id=study.id,
            file_path=file_path,
            respondent_label=respondent_label,
            settings=settings,
            embedding_client=embedding_client,
        )
        sources.append(source)

    session.flush()

    assert [source.source_code for source in sources] == ["S01", "S02", "S03"]
    assert all(source.status == "indexed" for source in sources)

    chunks = session.scalars(select(Chunk).order_by(Chunk.position)).all()
    assert len(chunks) >= 3
    for chunk in chunks:
        assert chunk.embedding is not None
        assert len(chunk.embedding) == settings.voyage_embedding_dimensions
        assert chunk.text.strip()

    for source in sources:
        source_chunks = session.scalars(select(Chunk).where(Chunk.source_id == source.id)).all()
        assert len(source_chunks) >= 1
