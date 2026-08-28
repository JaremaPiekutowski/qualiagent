"""Tests for DOCX report assembly."""

from io import BytesIO
from pathlib import Path

from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session

from qualiagent.config import Settings
from qualiagent.graph.run import run_main_path
from qualiagent.ingest.source_ingest import ingest_source
from qualiagent.models import AnalysisRun, Citation, Section, Study
from qualiagent.report import build_analysis_run_report, write_analysis_run_report
from tests.stub_embedding import StubEmbeddingClient
from tests.stub_language_model import StubLanguageModelClient
from tests.test_source_ingest import write_sample_docx, write_sample_pdf, write_sample_txt


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_build_analysis_run_report_contains_chapters_table_and_methodology(
    session: Session,
    settings: Settings,
    embedding_client: StubEmbeddingClient,
    tmp_path: Path,
) -> None:
    study = Study(
        name="Raport study",
        research_questions=["Jak respondenci postrzegają zmianę organizacyjną?"],
        web_search_enabled=False,
    )
    session.add(study)
    session.flush()

    for file_path, respondent_label in [
        (write_sample_txt(tmp_path), "R01"),
        (write_sample_docx(tmp_path), "R02"),
        (write_sample_pdf(tmp_path), "R03"),
    ]:
        ingest_source(
            session=session,
            study_id=study.id,
            file_path=file_path,
            respondent_label=respondent_label,
            settings=settings,
            embedding_client=embedding_client,
        )
    session.flush()

    final_state = run_main_path(
        session=session,
        study_id=study.id,
        settings=settings,
        embedding_client=embedding_client,
        language_model=StubLanguageModelClient(),
    )

    assert final_state["report_path"]
    report_path = Path(final_state["report_path"])
    assert report_path.exists()
    assert report_path.suffix == ".docx"

    analysis_run = session.execute(select(AnalysisRun).where(AnalysisRun.study_id == study.id)).scalar_one()
    report_bytes = build_analysis_run_report(session, analysis_run.id)
    document = Document(BytesIO(report_bytes))
    text = _document_text(document)

    assert study.name in text
    assert "Rozdziały" in text
    assert study.research_questions[0] in text
    assert "Tabela cytowań" in text
    assert "Marker" in text
    assert "Nota metodologiczna" in text
    assert "Liczba źródeł: 3." in text
    assert "Liczba respondentów: 3." in text
    assert "Liczba cytatów niezweryfikowanych: 0." in text
    assert document.tables
    assert any(cell.text.startswith("[S") for row in document.tables[0].rows for cell in row.cells)


def test_write_analysis_run_report_marks_unverified_citations(
    session: Session,
    settings: Settings,
    tmp_path: Path,
) -> None:
    study = Study(
        name="Unverified study",
        research_questions=["Pytanie cienkie?"],
        web_search_enabled=False,
    )
    session.add(study)
    session.flush()

    analysis_run = AnalysisRun(study_id=study.id, thread_id="thread-report-1", status="completed")
    session.add(analysis_run)
    session.flush()

    section = Section(
        analysis_run_id=analysis_run.id,
        research_question="Pytanie cienkie?",
        position=0,
        body="Tekst bez solidnego pokrycia.",
        coverage="thin",
        coverage_note="Za mało głosów.",
        respondents_covered=1,
        respondents_total=1,
    )
    session.add(section)
    session.flush()

    # Citation without a real source/chunk still documents verification flag in the table.
    # Use the study itself is invalid FK — create a minimal source+chunk via ingest instead.
    source_path = write_sample_txt(tmp_path)
    ingested = ingest_source(
        session=session,
        study_id=study.id,
        file_path=source_path,
        respondent_label="R01",
        settings=settings,
        embedding_client=StubEmbeddingClient(dimensions=settings.voyage_embedding_dimensions),
    )
    chunk = ingested.chunks[0]
    session.add(
        Citation(
            section_id=section.id,
            marker=f"[{ingested.source_code}:c{chunk.position}]",
            source_id=ingested.id,
            chunk_id=chunk.id,
            quoted_text="nieistniejący cytat",
            verified=False,
            verification_note="failed",
        )
    )
    session.flush()

    output_path = tmp_path / "reports" / "manual.docx"
    write_analysis_run_report(session, analysis_run.id, output_path)
    document = Document(str(output_path))
    text = _document_text(document)
    assert "Pytania z cienkim pokryciem:" in text
    assert "Pytanie cienkie?" in text
    assert "niezweryfikowany" in text
    assert "Liczba cytatów niezweryfikowanych: 1." in text
